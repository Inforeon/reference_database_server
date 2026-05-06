from __future__ import annotations

from typing import Any

from docx import Document as DocxDocument

from docsearch.extractors.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extractor for DOCX documents using python-docx."""

    @property
    def supported_extensions(self) -> list[str]:
        """Return list of file extensions this extractor handles."""
        return ["docx"]

    def extract_metadata(self, filepath: str) -> dict[str, Any]:
        """Extract structured metadata from a DOCX file."""
        try:
            doc = DocxDocument(filepath)
            cp = doc.core_properties
            result = {}
            for key in ("title", "author", "subject", "keywords", "comments"):
                value = getattr(cp, key, None)
                if value:
                    result[key] = value
            for key in ("created", "modified"):
                value = getattr(cp, key, None)
                if value is not None:
                    result[key] = value.isoformat()
            return result
        except Exception:
            return {}

    def extract_text(self, filepath: str) -> str:
        """Extract full text content from a DOCX file."""
        try:
            doc = DocxDocument(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception:
            return ""
